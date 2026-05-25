# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', size=10.5, bold=False):
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    if level == 1:
        set_chinese_font(run, '黑体', 16, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        set_chinese_font(run, '黑体', 14, bold=True)
    else:
        set_chinese_font(run, '黑体', 12, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_normal_text(doc, text):
    """添加普通正文"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_chinese_font(run, '宋体', 10.5)
    return p

# 创建文档
doc = Document()

# 设置默认中文字体
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(10.5)

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('论文阅读报告：SQIsign2D2数字签名方案')
set_chinese_font(title_run, '黑体', 18, bold=True)
title.paragraph_format.space_after = Pt(12)

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('——基于一维光滑同源的SQIsign2D优化变种')
set_chinese_font(subtitle_run, '楷体', 12)
subtitle.paragraph_format.space_after = Pt(24)

# 添加分隔线
doc.add_paragraph('─' * 50)

# 一、引言
add_heading_zh(doc, '一、引言', 1)

add_heading_zh(doc, '1.1 研究背景', 2)

add_normal_text(doc, 
    '2024年8月，NIST（美国国家标准与技术研究院）正式发布了首批后量子密码标准，'
    '这标志着密码学进入了后量子时代。在后量子密码的几种技术路线中，'
    '基于同源的密码学（Isogeny-based Cryptography）因其独特的数学结构而备受关注。')

add_normal_text(doc,
    '同源密码学的核心优势在于：目前尚无已知的量子算法能在多项式时间内求解超奇异椭圆曲线同源问题。'
    '这与基于格的方案不同——格密码虽然能抵抗已知的量子攻击，但其安全性依赖于一些启发式假设。'
    '同源问题的数学基础更为经典，与我们所学的椭圆曲线密码学（ECC）有很深的渊源，但又能够天然抵抗Shor算法的攻击。')

add_heading_zh(doc, '1.2 SQIsign与同源数字签名', 2)

add_normal_text(doc,
    'SQIsign是由De Feo等人于2020年提出的同源数字签名方案，也是目前唯一进入NIST第二轮评估的同源签名方案。'
    '它的最大优势是签名长度极短（约200字节，与ECC相当），但缺点是计算效率较低——'
    '签名和验证过程需要进行复杂的理想-同源转换。'
    '为了解决这一瓶颈，近年来陆续提出了SQIsignHD、SQIsign2D-West/East等改进方案，'
    '而本文介绍的SQIsign2D2则是最新的优化变种。')

add_heading_zh(doc, '1.3 本文核心贡献', 2)

add_normal_text(doc,
    '本文提出的SQIsign2D2通过重新设计素数参数和创新算法，实现了目前最快的SQIsign2D变体。'
    '在NIST-I安全级别下，与SQIsign2D-East相比：')

# 添加表格
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'

# 表头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '操作阶段'
hdr_cells[1].text = 'SQIsign2D-East'
hdr_cells[2].text = 'SQIsign2D2 (本文)'

# 数据行
row1 = table.rows[1].cells
row1[0].text = '密钥生成'
row1[1].text = '560 ms'
row1[2].text = '213 ms (快2.6倍)'

row2 = table.rows[2].cells
row2[0].text = '签名'
row2[1].text = '1263 ms'
row2[2].text = '587 ms (快2.1倍)'

row3 = table.rows[3].cells
row3[0].text = '验证'
row3[1].text = '296 ms'
row3[2].text = '247 ms (快1.2倍)'

# 设置表格字体
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '宋体', 10.5)

doc.add_paragraph()  # 空行

add_normal_text(doc,
    '关键创新在于利用一维光滑同源（3的幂次）来降低二维同源的计算开销，'
    '体现了降维优化的算法设计思想。')

# 二、技术背景
add_heading_zh(doc, '二、技术背景', 1)

add_heading_zh(doc, '2.1 椭圆曲线同源基础', 2)

add_normal_text(doc,
    '在量子信息科学的学习中，我们已经接触过椭圆曲线密码学。同源（Isogeny）是椭圆曲线之间的特殊映射，'
    '保持群结构的同时将一条曲线映射到另一条曲线。对于超奇异椭圆曲线，其自同态环与四元数代数中的最大序存在一一对应关系，'
    '这就是Deuring对应。')

add_normal_text(doc,
    '同源密码学的核心计算瓶颈是理想-同源转换：需要将四元数理想中的元素显式计算为曲线上的有理映射。'
    '这是SQIsign家族效率低下的主要原因。')

add_heading_zh(doc, '2.2 高维同源与Kani引理', 2)

add_normal_text(doc,
    'Kani引理是近年来同源密码学的重要工具。简单来说，它允许我们通过构造二维同源（即阿贝尔簇之间的同源）'
    '来分解和计算一维同源。SQIsign2D家族的签名方案都基于这一思想：'
    '通过二维同源计算来避免直接进行昂贵的理想-同源转换。')

add_heading_zh(doc, '2.3 SQIsign2D-East简介', 2)

add_normal_text(doc,
    '作为对比基准，SQIsign2D-East使用形如 p = 2^{a+b} * f - 1 的素数，其中f是小因子。'
    '这种参数设置使得2的幂次挠子群可以被高效利用，但主要计算仍依赖于二维(2^e, 2^e)-同源，计算成本较高。')

# 三、SQIsign2D2的核心创新
add_heading_zh(doc, '三、SQIsign2D2的核心创新', 1)

add_heading_zh(doc, '3.1 参数重构：p = C * D - 1', 2)

add_normal_text(doc,
    '本文的核心创新是改变了素数的选取方式。传统方案使用 p = f * 2^a - 1，仅利用2的幂次。'
    '而SQIsign2D2采用：')

add_normal_text(doc, '    p = 2^e2 * 3^e3 - 1')

add_normal_text(doc,
    '其中 C = 3^e3 ≈ D = 2^e2 ≈ √p。这样设计的妙处在于：')

add_normal_text(doc,
    '1. E₀[C] 和 E₀[D] 都是有理挠子群（定义在基域上）')
add_normal_text(doc,
    '2. C-同源和D-同源都可以用Velu公式高效计算（一维）')
add_normal_text(doc,
    '3. 二维同源的度数需求从约p降至约√p')

add_heading_zh(doc, '3.2 算法一：ImRanIso（改进随机同源生成）', 2)

add_normal_text(doc,
    '这是用于密钥生成的核心算法。其基本思路是：先构造一个度数为 d(D/D\'-d) * C * D\' 的内自同态，'
    '然后通过二维同源分解得到目标d-同源。相比前人的工作，该算法巧妙地利用CD\'-同源'
    '（其中D\' = 2^{floor(e₂/2)} ≈ p^{1/4}）来降低二维同源的度数需求。')

add_normal_text(doc,
    '直观理解：通过增加一维同源的计算量（约3/4 * log₂p），大幅降低二维同源的复杂度（降至约1/2 * log₂p），'
    '从而实现整体加速。')

add_heading_zh(doc, '3.3 算法二：GenImRanIso（广义改进随机同源）', 2)

add_normal_text(doc,
    '这是用于签名阶段生成辅助同源的算法。挑战在于：此时我们从曲线Eₐ出发，其自同态环未知。'
    '本文的创新是利用Eichler序（Eichler order）来构造所需的内自同态，而非使用完整的最大序。')

add_normal_text(doc,
    '关键优势：完全避免了East方案中需要的(2^e, 2^e)-二维同源，转而使用更易计算的(D,D)-二维同源。')

add_heading_zh(doc, '3.4 响应同源条件的简化', 2)

add_normal_text(doc,
    'East方案要求响应同源的度数q满足(2^a, 2^b)3-nice的复杂条件。本文简化为仅需D-adequate：'
    'q为奇数、q < D、且3不整除q。简化后采样成功率从约1/8提升至约1/4。')

# 四、协议流程
add_heading_zh(doc, '四、SQIsign2D2协议流程', 1)

add_normal_text(doc,
    'SQIsign2D2是一个Σ-协议（识别协议），可通过Fiat-Shamir变换转换为数字签名方案。'
    '协议包含以下阶段：')

add_heading_zh(doc, '4.1 参数设置', 2)

add_normal_text(doc,
    '以NIST-I安全级别为例，选取 p = 2^131 * 3^78 - 1，这是一个约256比特的素数。'
    '公钥大小约64字节，签名大小约154字节（压缩版）。')

add_heading_zh(doc, '4.2 密钥生成', 2)

add_normal_text(doc,
    '1. 选取随机素数 N_τ < p^{1/4}，满足Legendre符号(3/N_τ) = -1（用于抵抗密钥恢复攻击）')
add_normal_text(doc,
    '2. 使用ImRanIso算法生成N_τ-同源 τ: E₀ → Eₐ')
add_normal_text(doc,
    '3. 公钥为曲线Eₐ，私钥包含同源τ的信息')

add_heading_zh(doc, '4.3 承诺-挑战-响应', 2)

add_normal_text(doc,
    '承诺：证明者使用ImRanIso生成随机同源 ψ: E₀ → E₁，发送承诺曲线E₁。')

add_normal_text(doc,
    '挑战：验证者发送随机点 K_cha ∈ E₁[C] 作为挑战。')

add_normal_text(doc,
    '响应：证明者计算挑战同源 φ: E₁ → E₂，然后构造响应同源 σ: Eₐ → E₂。'
    '利用GenImRanIso生成辅助同源，最终返回压缩后的响应。')

add_heading_zh(doc, '4.4 验证', 2)

add_normal_text(doc,
    '验证者通过计算(D,D)-二维同源来验证响应的有效性。若同源分解后得到的σ满足条件，则接受。')

# 五、安全性分析
add_heading_zh(doc, '五、安全性分析', 1)

add_heading_zh(doc, '5.1 特殊可靠性', 2)

add_normal_text(doc,
    '若证明者能对两个不同的挑战生成有效响应，则可提取出Eₐ的非平凡自同态，'
    '进而可能计算出密钥。这保证了协议的可靠性。')

add_heading_zh(doc, '5.2 零知识性', 2)

add_normal_text(doc,
    '在随机预言机模型下，协议满足特殊诚实验证者零知识。核心假设是：'
    'ImRanIso输出的承诺曲线与随机超奇异曲线计算不可区分。')

add_heading_zh(doc, '5.3 承诺采样的强化', 2)

add_normal_text(doc,
    '针对上述启发式假设，论文提出了两种强化方案：'
    '（1）Double Path技术：将承诺同源度数提升至约p，增强随机性；'
    '（2）(3,3)-同源适配：直接计算(CD,CD)-二维同源。')

# 六、性能评估
add_heading_zh(doc, '六、性能评估与对比', 1)

add_normal_text(doc,
    '实验在Intel Core i9-12900K上使用Julia实现。核心发现：')

add_normal_text(doc,
    '• SQIsign2D2用一维3-同源替代了大量二维(2,2)-同源，这是效率提升的关键')
add_normal_text(doc,
    '• 密钥生成阶段，(2,2)-同源计算从253步降至66步')
add_normal_text(doc,
    '• 签名阶段，(2,2)-同源计算从641步降至262步')

add_normal_text(doc,
    '与同期工作SQIsign2DPush相比，本文方案在承诺生成阶段依赖更强的安全假设，但效率更高。'
    '与PRISM-sig相比，本文方案验证更快。')

# 七、个人思考
add_heading_zh(doc, '七、个人思考与展望', 1)

add_heading_zh(doc, '7.1 量子信息视角的解读', 2)

add_normal_text(doc,
    '作为量子信息专业的学生，我对同源密码学的量子安全性有一些思考。Shor算法能破解RSA/ECC，'
    '是因为它们依赖于阿贝尔群上的隐藏子群问题。然而，超奇异同源问题对应的是非阿贝尔群'
    '（四元数代数理想类群），目前尚无有效的量子算法。')

add_normal_text(doc,
    '本文通过增加一维光滑同源的使用来优化性能，这一策略的量子安全性很有意思：'
    '一维光滑同源的经典计算已经高效（Velu公式），量子计算也没有明显优势；'
    '而二维同源即使对量子计算机也是难题。这种经典-量子的权衡设计值得深入研究。')

add_heading_zh(doc, '7.2 对论文的评价', 2)

add_normal_text(doc, '优点：')
add_normal_text(doc,
    '1. 问题定位准确：作者精准识别出二维同源是效率瓶颈，而不是盲目优化')
add_normal_text(doc,
    '2. 工程实现完整：提供了Julia语言的完整实现，实验数据详实')
add_normal_text(doc,
    '3. 理论-实践平衡：在安全性假设与效率之间取得了较好的权衡')

add_normal_text(doc, '可改进之处：')
add_normal_text(doc,
    '1. 承诺生成的随机性假设缺乏严格证明，虽然有强化方案但会增加开销')
add_normal_text(doc,
    '2. 即使优化后，签名时间（587ms）仍远高于NIST标准的ML-DSA（<1ms）'
    '——虽然签名长度优势明显（154B vs 2420B）')

add_heading_zh(doc, '7.3 未来展望', 2)

add_normal_text(doc,
    '短期来看，可以进一步优化(3,3)-同源的实现，使安全强化方案更实用；'
    '探索GPU/FPGA加速，特别是KaniCod算法的并行化。')

add_normal_text(doc,
    '中期来看，值得研究同源密码与其他后量子方案（如格密码）的混合协议，'
    '实现双重保护。此外，后量子TLS/SSH的集成也是重要的工程问题。')

add_normal_text(doc,
    '长期来看，当容错量子计算真正实现时，同源密码是否仍安全？'
    '这需要量子复杂性理论的突破，特别是非阿贝尔隐藏子群问题的量子难度证明。')

# 八、结论
add_heading_zh(doc, '八、结论', 1)

add_normal_text(doc,
    'SQIsign2D2通过参数重构（p = 2^e2 * 3^e3 - 1）和算法创新（ImRanIso/GenImRanIso），'
    '将SQIsign2D家族的效率提升到了新高度。其核心洞见——用一维光滑同源替代高维同源计算——'
    '不仅具有直接的工程价值，也为后量子密码学的经典-量子资源权衡提供了新思路。')

add_normal_text(doc,
    '作为量子信息专业的学生，我认为本文最重要的启示是：量子安全不等于量子低效。'
    '通过精细的算法设计，完全可以在保持抗量子安全性的同时，实现可实用的计算效率。')

# 参考文献
doc.add_page_break()
add_heading_zh(doc, '参考文献', 1)

refs = [
    '[1] NIST. Post-Quantum Cryptography Standardization. 2024.',
    '[2] De Feo L, et al. SQIsign: Compact Post-Quantum Signatures from Quaternions and Isogenies. ASIACRYPT 2020.',
    '[3] Dartois P, et al. SQIsignHD: New Dimensions in Cryptography. EUROCRYPT 2024.',
    '[4] Nakagawa K, et al. SQIsign2D-East: A New Signature Scheme Using 2-Dimensional Isogenies. ASIACRYPT 2024.',
    '[5] Xu Z, et al. SQIsign2D2: New SQIsign2D Variant by Leveraging Power Smooth Isogenies in Dimension One. 2025.',
    '[6] Castryck W, et al. Breaking and Repairing SQIsign2D-East. ePrint 2024/1453.',
    '[7] Kani E. The number of curves of genus two with elliptic differentials. J. Reine Angew. Math. 1997.'
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(ref)
    set_chinese_font(run, '宋体', 10.5)

# 保存文档
doc.save('阅读报告_修订版.docx')
print('Report generated: 阅读报告_修订版.docx')
